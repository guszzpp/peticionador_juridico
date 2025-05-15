from docx import Document
from pathlib import Path
from typing import Dict


def substituir_placeholders_em_docx(
    caminho_modelo: Path,
    caminho_saida: Path,
    substituicoes: Dict[str, str]
) -> None:
    """
    Substitui os placeholders no modelo .docx e salva o resultado no caminho de saída.
    Preserva formatação (negrito, itálico, alinhamento etc.) do documento original.

    Args:
        caminho_modelo (Path): Caminho do arquivo .docx com placeholders no formato {{CHAVE}}.
        caminho_saida (Path): Caminho onde o novo .docx formatado será salvo.
        substituicoes (Dict[str, str]): Dicionário com os placeholders (sem chaves duplas) e seus valores.
    """
    # Garantir que todos os valores no dicionário de substituição sejam strings
    substituicoes_str = {
        k: str(v) if v is not None else "" 
        for k, v in substituicoes.items()
    }
    
    doc = Document(caminho_modelo)

    # Substitui nos parágrafos principais
    for paragrafo in doc.paragraphs:
        _substituir_em_paragrafo(paragrafo, substituicoes_str)

    # Substitui em células de tabelas, se houver
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    _substituir_em_paragrafo(paragrafo, substituicoes_str)

    doc.save(caminho_saida)


def _substituir_em_paragrafo(paragrafo, substituicoes: Dict[str, str]) -> None:
    """
    Substitui os placeholders dentro de um parágrafo, run por run,
    mantendo a formatação original.
    """
    for chave, valor in substituicoes.items():
        placeholder = f"{{{{{chave}}}}}"  # Ex: {{NUM_PROCESSO}}
        if placeholder in paragrafo.text:
            for run in paragrafo.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, valor)