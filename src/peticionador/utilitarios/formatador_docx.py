import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from pathlib import Path
from typing import Dict

def criar_docx_formatado(texto_completo: str, caminho_saida: Path, dados_documento: Dict[str, str] = None) -> None:
    """
    Cria um documento DOCX bem formatado a partir do texto fornecido.
    
    Args:
        texto_completo (str): Texto completo da minuta
        caminho_saida (Path): Caminho para salvar o arquivo .docx
        dados_documento (Dict[str, str], optional): Dados adicionais como título, tipo de recurso, etc.
    """
    if dados_documento is None:
        dados_documento = {}
    
    doc = Document()
    
    # Configurar página
    sections = doc.sections
    for section in sections:
        section.page_height = Cm(29.7)  # A4
        section.page_width = Cm(21.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
    
    # Definir estilos
    styles = doc.styles
    
    # Estilo para título principal 
    titulo_style = styles.add_style('TituloPrincipal', WD_STYLE_TYPE.PARAGRAPH)
    titulo_font = titulo_style.font
    titulo_font.name = 'Arial'
    titulo_font.size = Pt(14)
    titulo_font.bold = True
    titulo_paragraph_format = titulo_style.paragraph_format
    titulo_paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo_paragraph_format.space_after = Pt(12)
    
    # Estilo para subtítulos (teses)
    subtitulo_style = styles.add_style('Subtitulo', WD_STYLE_TYPE.PARAGRAPH)
    subtitulo_font = subtitulo_style.font
    subtitulo_font.name = 'Arial'
    subtitulo_font.size = Pt(12)
    subtitulo_font.bold = True
    subtitulo_paragraph_format = subtitulo_style.paragraph_format
    subtitulo_paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitulo_paragraph_format.space_before = Pt(12)
    subtitulo_paragraph_format.space_after = Pt(6)
    
    # Estilo para texto normal
    texto_style = styles.add_style('TextoNormal', WD_STYLE_TYPE.PARAGRAPH)
    texto_font = texto_style.font
    texto_font.name = 'Arial'
    texto_font.size = Pt(12)
    texto_paragraph_format = texto_style.paragraph_format
    texto_paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    texto_paragraph_format.space_after = Pt(6)
    texto_paragraph_format.line_spacing = 1.15
    
    # Adicionar título principal
    tipo_recurso = dados_documento.get('tipo_recurso', 'RECURSO')
    titulo_texto = f"CONTRARRAZÕES AO {tipo_recurso.upper()}"
    titulo = doc.add_paragraph(titulo_texto, style='TituloPrincipal')
    
    # Processar o texto completo
    paragrafos = texto_completo.strip().split('\n\n')
    
    for paragrafo in paragrafos:
        paragrafo = paragrafo.strip()
        
        # Verificar se é um subtítulo/tese (geralmente começa com número e ponto)
        if re.match(r'^\d+\.\s+[A-Z]', paragrafo):
            p = doc.add_paragraph(paragrafo, style='Subtitulo')
        else:
            p = doc.add_paragraph(paragrafo, style='TextoNormal')
    
    # Salvar o documento
    doc.save(str(caminho_saida))
    
    return True


def criar_odt_formatado(texto_completo: str, caminho_saida: Path, dados_documento: Dict[str, str] = None) -> None:
    """
    Cria um documento ODT bem formatado a partir do texto fornecido.
    
    Args:
        texto_completo (str): Texto completo da minuta
        caminho_saida (Path): Caminho para salvar o arquivo .odt
        dados_documento (Dict[str, str], optional): Dados adicionais como título, tipo de recurso, etc.
    """
    try:
        from odf.opendocument import OpenDocumentText
        from odf.style import Style, TextProperties, ParagraphProperties, TabStop, TabStops, FontFace
        from odf.text import P, H, Span
        
        if dados_documento is None:
            dados_documento = {}
        
        # Criar documento
        doc_odt = OpenDocumentText()
        
        # Registrar fonte
        doc_odt.fontfacedecls.addElement(FontFace(name="Arial", fontfamily="Arial", fontfamilygeneric="swiss"))
        
        # Definir estilos
        # Estilo para título
        titulo_style = Style(name="TituloStyle", family="paragraph")
        titulo_style.addElement(TextProperties(fontname="Arial", fontweight="bold", fontsize="14pt"))
        titulo_style.addElement(ParagraphProperties(textalign="center", marginbottom="12pt"))
        doc_odt.styles.addElement(titulo_style)
        
        # Estilo para subtítulos
        subtitulo_style = Style(name="SubtituloStyle", family="paragraph")
        subtitulo_style.addElement(TextProperties(fontname="Arial", fontweight="bold", fontsize="12pt"))
        subtitulo_style.addElement(ParagraphProperties(textalign="left", margintop="12pt", marginbottom="6pt"))
        doc_odt.styles.addElement(subtitulo_style)
        
        # Estilo para texto normal justificado
        texto_style = Style(name="TextoStyle", family="paragraph")
        texto_style.addElement(TextProperties(fontname="Arial", fontsize="12pt"))
        texto_style.addElement(ParagraphProperties(textalign="justify", marginbottom="6pt", linespacing="115%"))
        doc_odt.styles.addElement(texto_style)
        
        # Adicionar título
        tipo_recurso = dados_documento.get('tipo_recurso', 'RECURSO')
        titulo_texto = f"CONTRARRAZÕES AO {tipo_recurso.upper()}"
        titulo = H(outlinelevel=1, stylename=titulo_style)
        titulo.addText(titulo_texto)
        doc_odt.text.addElement(titulo)
        
        # Processar o texto completo
        import re
        paragrafos = texto_completo.strip().split('\n\n')
        
        for paragrafo in paragrafos:
            paragrafo = paragrafo.strip()
            
            # Verificar se é um subtítulo/tese (geralmente começa com número e ponto)
            if re.match(r'^\d+\.\s+[A-Z]', paragrafo):
                p = P(stylename=subtitulo_style)
                p.addText(paragrafo)
                doc_odt.text.addElement(p)
            else:
                p = P(stylename=texto_style)
                p.addText(paragrafo)
                doc_odt.text.addElement(p)
        
        # Salvar documento
        doc_odt.save(str(caminho_saida))
        return True
        
    except Exception as e:
        print(f"Erro ao criar ODT formatado: {e}")
        return False