
from io import BytesIO
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from odf.opendocument import OpenDocumentText
from odf.style import Style, TextProperties, ParagraphProperties, FontFace
from odf.text import P, H, Span

def gerar_docx_formatado(texto_html: str) -> BytesIO:
    soup = BeautifulSoup(texto_html, "html.parser")
    doc = Document()

    for tag in soup.find_all(['p', 'div', 'hr']):
        texto = tag.get_text(strip=True)
        if not texto and tag.name != 'hr':
            continue

        if tag.name == 'hr':
            doc.add_paragraph().add_run().add_break()
            continue

        par = doc.add_paragraph()
        run = par.add_run(texto)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        par.paragraph_format.first_line_indent = Cm(1.5)
        par.paragraph_format.line_spacing = 1.5
        par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def gerar_odt_formatado(texto_html: str) -> BytesIO:
    soup = BeautifulSoup(texto_html, "html.parser")
    doc = OpenDocumentText()

    doc.fontfacedecls.addElement(FontFace(name="Times New Roman", fontfamily="Times New Roman", fontfamilygeneric="roman"))

    estilo_justificado = Style(name="Justificado", family="paragraph")
    estilo_justificado.addElement(TextProperties(fontname="Times New Roman", fontsize="12pt"))
    estilo_justificado.addElement(ParagraphProperties(
        textalign="justify",
        marginleft="1.5cm",
        marginright="0cm",
        margintop="0.2cm",
        marginbottom="0.2cm",
        lineheight="150%"
    ))
    doc.styles.addElement(estilo_justificado)

    for tag in soup.find_all(['p', 'div', 'hr']):
        texto = tag.get_text(strip=True)
        if not texto and tag.name != 'hr':
            continue

        if tag.name == 'hr':
            p = P(stylename=estilo_justificado)
            p.addText("------------------------------")
            doc.text.addElement(p)
            continue

        p = P(stylename=estilo_justificado)
        p.addText(texto)
        doc.text.addElement(p)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
